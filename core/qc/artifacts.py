"""Persistence and export helpers for immutable QC reports."""
from __future__ import annotations

import json
from dataclasses import dataclass
from pathlib import Path
from tempfile import NamedTemporaryFile
from typing import Callable

from core.build_info import BUILD_VERSION
from core.data_platform import (
    ArtifactStore,
    DatasetManifest,
    DatasetManifestRepository,
    DatasetMetadataCatalog,
    DatasetProvenance,
)

from .models import QCReport


@dataclass(frozen=True, slots=True)
class QCPersistenceResult:
    manifest: DatasetManifest
    report_relative_path: str

    def to_dict(self) -> dict[str, object]:
        return {
            "dataset_id": self.manifest.dataset_id,
            "lineage_id": self.manifest.lineage_id,
            "version": self.manifest.version,
            "source_dataset_ids": list(self.manifest.provenance.source_dataset_ids),
            "artifact_path": self.report_relative_path,
            "checksum_sha256": self.manifest.checksum_sha256,
            "size_bytes": self.manifest.size_bytes,
        }




@dataclass(frozen=True, slots=True)
class QCExportResult:
    manifest: DatasetManifest
    export_relative_path: str
    format_id: str

    def to_dict(self) -> dict[str, object]:
        return {
            "dataset_id": self.manifest.dataset_id,
            "lineage_id": self.manifest.lineage_id,
            "version": self.manifest.version,
            "source_dataset_ids": list(self.manifest.provenance.source_dataset_ids),
            "artifact_path": self.export_relative_path,
            "format_id": self.format_id,
            "checksum_sha256": self.manifest.checksum_sha256,
            "size_bytes": self.manifest.size_bytes,
        }


class QCReportArtifactService:
    """Store QC results without mutating the source Dataset lineage."""

    def __init__(self, projects_root: Path | str) -> None:
        self.projects_root = Path(projects_root)
        self.artifacts = ArtifactStore(self.projects_root)
        self.manifests = DatasetManifestRepository(self.projects_root)
        self.catalog = DatasetMetadataCatalog(self.projects_root)

    def persist(
        self,
        *,
        project_id: str,
        source_dataset_id: str,
        report: QCReport,
        actor: str = "",
    ) -> QCPersistenceResult:
        source = self.manifests.load(project_id, source_dataset_id)
        payload = report.to_dict()
        payload["source_dataset_id"] = source.dataset_id
        payload["source_lineage_id"] = source.lineage_id
        payload["source_version"] = source.version

        with NamedTemporaryFile("w", encoding="utf-8", suffix=".qc.json", delete=False) as handle:
            temp_path = Path(handle.name)
            json.dump(payload, handle, ensure_ascii=False, indent=2, sort_keys=True)
            handle.write("\n")
        try:
            filename = f"{source.dataset_id}-qc-{report.generated_at.replace(':', '').replace('+', '_')}.json"
            location = self.artifacts.store_file(
                project_id=project_id,
                source=temp_path,
                kind="derived",
                filename=filename,
            )
        finally:
            temp_path.unlink(missing_ok=True)

        manifest = DatasetManifest.create(
            project_id=project_id,
            well_id=source.well_id,
            format_id="qc-report-json",
            artifact_path=location.relative_path,
            checksum_sha256=location.checksum_sha256,
            size_bytes=location.size_bytes,
            source_name=filename,
            metadata={
                "dataset_kind": report.dataset_kind,
                "qc_status": report.status,
                "finding_count": len(report.findings),
                "error_count": sum(1 for item in report.findings if item.severity == "error"),
                "warning_count": sum(1 for item in report.findings if item.severity == "warning"),
                "source_dataset_id": source.dataset_id,
                "source_lineage_id": source.lineage_id,
                "source_version": source.version,
            },
            provenance=DatasetProvenance(
                operation="quality-control",
                actor=actor,
                source_dataset_ids=(source.dataset_id,),
                application_version=BUILD_VERSION,
            ),
        )
        self.manifests.save(manifest)
        self.catalog.project(manifest)
        return QCPersistenceResult(manifest=manifest, report_relative_path=location.relative_path)

    def export_docx(self, *, report: QCReport, destination: Path | str, translate: Callable[[str], str]) -> Path:
        try:
            from docx import Document
        except ImportError as exc:  # pragma: no cover - dependency environment specific
            raise RuntimeError("DOCX export requires python-docx") from exc
        path = Path(destination)
        path.parent.mkdir(parents=True, exist_ok=True)
        doc = Document()
        doc.add_heading(translate("qc.report.title"), level=0)
        doc.add_paragraph(f"{translate('qc.report.status')}: {translate(f'qc.status.{report.status}')}")
        doc.add_paragraph(f"{translate('qc.report.rows')}: {report.row_count}")
        doc.add_paragraph(f"{translate('qc.report.curves')}: {report.curve_count}")
        doc.add_heading(translate("qc.report.findings"), level=1)
        for finding in report.findings:
            text = translate(finding.message_key)
            doc.add_paragraph(f"[{finding.severity.upper()}] {finding.code}: {text}", style="List Bullet")
        doc.add_heading(translate("qc.report.statistics"), level=1)
        table = doc.add_table(rows=1, cols=6)
        headers = ("curve", "valid", "null", "min", "max", "mean")
        for cell, key in zip(table.rows[0].cells, headers):
            cell.text = translate(f"qc.report.column.{key}")
        for item in report.curve_statistics:
            cells = table.add_row().cells
            values = (item.curve, item.valid_count, item.null_count, item.minimum, item.maximum, item.mean)
            for cell, value in zip(cells, values):
                cell.text = "" if value is None else str(value)
        doc.save(path)
        return path


    def export_and_register(
        self,
        *,
        project_id: str,
        source_qc_dataset_id: str,
        report: QCReport,
        format_id: str,
        translate: Callable[[str], str],
        actor: str = "",
    ) -> QCExportResult:
        """Create a report file and register it as an immutable export Dataset."""
        normalized = str(format_id).strip().lower()
        if normalized not in {"pdf", "docx"}:
            raise ValueError("QC export format must be pdf or docx")
        source = self.manifests.load(project_id, source_qc_dataset_id)
        suffix = f".{normalized}"
        with NamedTemporaryFile(suffix=suffix, delete=False) as handle:
            temp_path = Path(handle.name)
        try:
            if normalized == "pdf":
                self.export_pdf(report=report, destination=temp_path, translate=translate)
            else:
                self.export_docx(report=report, destination=temp_path, translate=translate)
            filename = f"{source.dataset_id}-qc-report{suffix}"
            location = self.artifacts.store_file(
                project_id=project_id, source=temp_path, kind="exports", filename=filename
            )
        finally:
            temp_path.unlink(missing_ok=True)

        manifest = DatasetManifest.create(
            project_id=project_id,
            well_id=source.well_id,
            format_id=normalized,
            artifact_path=location.relative_path,
            checksum_sha256=location.checksum_sha256,
            size_bytes=location.size_bytes,
            source_name=filename,
            metadata={
                "dataset_kind": "qc-report-export",
                "report_format": normalized,
                "qc_status": report.status,
                "source_qc_dataset_id": source.dataset_id,
            },
            provenance=DatasetProvenance(
                operation="quality-control-export",
                actor=actor,
                source_dataset_ids=(source.dataset_id,),
                application_version=BUILD_VERSION,
            ),
        )
        self.manifests.save(manifest)
        self.catalog.project(manifest)
        return QCExportResult(manifest, location.relative_path, normalized)

    def export_pdf(self, *, report: QCReport, destination: Path | str, translate: Callable[[str], str]) -> Path:
        try:
            from reportlab.lib.pagesizes import A4
            from reportlab.lib.styles import getSampleStyleSheet
            from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle
            from reportlab.lib import colors
        except ImportError as exc:  # pragma: no cover
            raise RuntimeError("PDF export requires reportlab") from exc
        path = Path(destination)
        path.parent.mkdir(parents=True, exist_ok=True)
        styles = getSampleStyleSheet()
        story = [
            Paragraph(translate("qc.report.title"), styles["Title"]),
            Paragraph(f"{translate('qc.report.status')}: {translate(f'qc.status.{report.status}')}", styles["BodyText"]),
            Spacer(1, 10),
            Paragraph(translate("qc.report.findings"), styles["Heading2"]),
        ]
        for finding in report.findings:
            story.append(Paragraph(f"[{finding.severity.upper()}] {finding.code}: {translate(finding.message_key)}", styles["BodyText"]))
        story.extend([Spacer(1, 10), Paragraph(translate("qc.report.statistics"), styles["Heading2"])])
        data = [[translate(f"qc.report.column.{key}") for key in ("curve", "valid", "null", "min", "max", "mean")]]
        for item in report.curve_statistics:
            data.append([item.curve, item.valid_count, item.null_count, item.minimum, item.maximum, item.mean])
        table = Table(data, repeatRows=1)
        table.setStyle(TableStyle([
            ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
            ("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ]))
        story.append(table)
        SimpleDocTemplate(str(path), pagesize=A4, title=translate("qc.report.title")).build(story)
        return path
