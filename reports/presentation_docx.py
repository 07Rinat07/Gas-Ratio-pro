from __future__ import annotations

from dataclasses import dataclass
from io import BytesIO
from pathlib import Path
from typing import Sequence, Callable

try:
    from docx import Document
    from docx.enum.section import WD_ORIENT
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt, RGBColor
    DOCX_AVAILABLE = True
except ModuleNotFoundError:  # pragma: no cover - depends on user environment
    Document = None
    WD_ORIENT = WD_TABLE_ALIGNMENT = WD_CELL_VERTICAL_ALIGNMENT = WD_ALIGN_PARAGRAPH = None
    Inches = Pt = RGBColor = None
    DOCX_AVAILABLE = False

from reports.document_model import (
    DocumentNotice,
    DocumentPlot,
    DocumentVisualizationPreview,
    DocumentTable,
    EngineeringDocument,
    build_engineering_document,
)
from reports.presentation_model import PresentationModel
from reports.plot_theme import apply_report_plot_theme


@dataclass(frozen=True)
class PresentationDocxOptions:
    """Options for renderer-neutral DOCX export.

    The DOCX renderer consumes EngineeringDocument only. It does not rerun gas
    ratio calculations, interpretation rules, evidence selection or report-card
    logic. This keeps HTML, PDF and DOCX synchronized through the same document
    model.
    """

    include_figures: bool = True
    include_technical_appendix: bool = False
    paper_size: str = "A4"
    orientation: str = "portrait"
    margin_mm: int = 14
    title: str = "Gas Ratio Professional Report"


@dataclass(frozen=True)
class PresentationDocxResult:
    """Rendered DOCX payload and document composition metadata."""

    content: bytes
    profile: str
    table_titles: tuple[str, ...]
    figure_count: int
    schema: str = "gas-ratio-pro/presentation/docx/v1"


def ensure_docx_available() -> None:
    """Raise a clear runtime error when DOCX export dependency is missing."""

    if not DOCX_AVAILABLE:
        raise RuntimeError(
            "DOCX export requires the optional dependency 'python-docx'. "
            "Install project dependencies with: pip install -r requirements.txt"
        )


def _clean_text(value: object) -> str:
    return str(value if value is not None else "").strip()


def _safe_orientation(value: str) -> str:
    text = str(value or "portrait").strip().lower()
    return text if text in {"portrait", "landscape"} else "portrait"


def _safe_margin_mm(value: int) -> int:
    try:
        number = int(value)
    except (TypeError, ValueError):
        number = 14
    return max(6, min(number, 25))


def _safe_paper_size(value: str) -> tuple[float, float]:
    # Dimensions are inches because python-docx exposes section sizes in EMUs
    # via Length helpers. The values are standard ISO/US page sizes.
    text = str(value or "A4").strip().upper()
    sizes = {
        "A4": (8.27, 11.69),
        "A3": (11.69, 16.54),
        "LETTER": (8.5, 11.0),
    }
    return sizes.get(text, sizes["A4"])


def _configure_document(doc: Document, options: PresentationDocxOptions) -> None:
    section = doc.sections[0]
    width, height = _safe_paper_size(options.paper_size)
    if _safe_orientation(options.orientation) == "landscape":
        width, height = height, width
        section.orientation = WD_ORIENT.LANDSCAPE
    else:
        section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(width)
    section.page_height = Inches(height)
    margin = Inches(_safe_margin_mm(options.margin_mm) / 25.4)
    section.left_margin = margin
    section.right_margin = margin
    section.top_margin = margin
    section.bottom_margin = margin

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10)
    for style_name, size in (("Title", 18), ("Heading 1", 14), ("Heading 2", 12)):
        try:
            styles[style_name].font.name = "Arial"
            styles[style_name].font.size = Pt(size)
        except KeyError:
            pass


def _add_paragraph(doc: Document, text: object, *, style: str | None = None, bold: bool = False) -> None:
    paragraph = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    run = paragraph.add_run(_clean_text(text))
    run.bold = bold


def _add_metadata_table(doc: Document, rows: Sequence[tuple[str, str]]) -> None:
    clean_rows = [(label, value) for label, value in rows if _clean_text(label) and _clean_text(value)]
    if not clean_rows:
        return
    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    for label, value in clean_rows:
        cells = table.add_row().cells
        cells[0].text = _clean_text(label)
        cells[1].text = _clean_text(value)
        for paragraph in cells[0].paragraphs:
            for run in paragraph.runs:
                run.bold = True
        for cell in cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    doc.add_paragraph()




def _adaptive_docx_column_widths(headers: Sequence[str], rows: Sequence[Sequence[str]], *, total_inches: float = 6.6) -> list[float]:
    if not headers:
        return []
    weights: list[float] = []
    for index, header in enumerate(headers):
        samples = [str(header)] + [str(row[index]) for row in rows[:40] if index < len(row)]
        longest = max((len(value.strip()) for value in samples), default=1)
        weights.append(max(4.0, min(22.0, longest ** 0.72)))
    total = sum(weights) or 1.0
    minimum = 0.55 if len(headers) >= 7 else 0.75
    raw = [total_inches * weight / total for weight in weights]
    adjusted = [max(minimum, width) for width in raw]
    scale = total_inches / sum(adjusted)
    return [width * scale for width in adjusted]


def _add_document_table(doc: Document, block: DocumentTable) -> None:
    if not block.headers or not block.rows:
        return
    _add_paragraph(doc, block.title or "Таблица", style="Heading 2")
    table = doc.add_table(rows=1, cols=len(block.headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    table.autofit = False
    column_widths = _adaptive_docx_column_widths(block.headers, block.rows)
    header_cells = table.rows[0].cells
    for index, header in enumerate(block.headers):
        header_cells[index].width = Inches(column_widths[index])
        header_cells[index].text = _clean_text(header)
        for paragraph in header_cells[index].paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(8.5)
    max_cols = len(block.headers)
    for source_row in block.rows:
        row_cells = table.add_row().cells
        cells = list(source_row[:max_cols])
        if len(cells) < max_cols:
            cells.extend([""] * (max_cols - len(cells)))
        for index, value in enumerate(cells):
            row_cells[index].width = Inches(column_widths[index])
            row_cells[index].text = _clean_text(value)
            row_cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            for paragraph in row_cells[index].paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(8.5)
    doc.add_paragraph()


def _add_notice(doc: Document, block: DocumentNotice) -> None:
    _add_paragraph(doc, block.title or "Примечание", style="Heading 2")
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.15)
    run = paragraph.add_run(_clean_text(block.text))
    run.italic = block.role not in {"error", "warning"}
    doc.add_paragraph()


def _add_visualization_preview_placeholder(doc: Document, block: DocumentVisualizationPreview) -> None:
    _add_paragraph(doc, block.title or "LAS visualization preview", style="Heading 2")
    preview = dict(block.preview or {})
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(
        f"SVG preview prepared by Visualization Engine: "
        f"tracks={preview.get('track_count', 0)}, curves={preview.get('curve_count', 0)}, overlays={preview.get('overlay_count', 0)}."
    )
    run.italic = True
    doc.add_paragraph()


def _figure_report_legend(figure: object) -> dict[str, object]:
    layout = getattr(figure, "layout", None)
    meta = getattr(layout, "meta", None) if layout is not None else None
    if not isinstance(meta, dict):
        return {}
    payload = meta.get("gas_ratio_report_legend", {})
    return dict(payload) if isinstance(payload, dict) else {}


def _hex_rgb(value: object) -> tuple[int, int, int]:
    text = str(value or "#64748b").strip().lstrip("#")
    if len(text) != 6:
        return (100, 116, 139)
    try:
        return tuple(int(text[index:index + 2], 16) for index in (0, 2, 4))
    except ValueError:
        return (100, 116, 139)


def _add_report_legend_table(doc: Document, title: str, entries: Sequence[dict[str, object]]) -> None:
    if not entries:
        return
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(title)
    run.bold = True
    run.font.size = Pt(11)
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    headers = ("Знак", "Обозначение", "Инженерное значение")
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.text = header
        for cell_run in cell.paragraphs[0].runs:
            cell_run.bold = True
            cell_run.font.size = Pt(10)
    for entry in entries:
        cells = table.add_row().cells
        symbol = str(entry.get("symbol", "■"))
        color = _hex_rgb(entry.get("color"))
        cells[0].text = ""
        symbol_run = cells[0].paragraphs[0].add_run(symbol)
        symbol_run.bold = True
        symbol_run.font.size = Pt(12)
        symbol_run.font.color.rgb = RGBColor(*color)
        cells[1].text = str(entry.get("label", ""))
        cells[2].text = str(entry.get("description", ""))
        for cell in cells[1:]:
            for cell_run in cell.paragraphs[0].runs:
                cell_run.font.size = Pt(10)
        for cell in cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    doc.add_paragraph()


def _add_statistics_table(doc: Document, entries: Sequence[dict[str, object]]) -> None:
    if not entries:
        return
    paragraph = doc.add_paragraph()
    run = paragraph.add_run("Статистика кривых")
    run.bold = True
    run.font.size = Pt(11)
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for index, header in enumerate(("Кривая", "Мин.", "Макс.", "Среднее", "Сумма")):
        table.rows[0].cells[index].text = header
    for entry in entries:
        cells = table.add_row().cells
        values = (str(entry.get("label", "")), f"{float(entry.get('minimum',0)):.4g}",
                  f"{float(entry.get('maximum',0)):.4g}", f"{float(entry.get('mean',0)):.4g}",
                  f"{float(entry.get('sum',0)):.5g}")
        for index, value in enumerate(values):
            cells[index].text = value
    doc.add_paragraph()


def _add_plot_placeholder(doc: Document, block: DocumentPlot) -> None:
    """Embed the canonical vector composite or a legacy Plotly figure into DOCX."""
    _add_paragraph(doc, block.title or "Планшет", style="Heading 2")
    if hasattr(block.figure, "svg"):
        from svglib.svglib import svg2rlg
        from reportlab.graphics import renderPM
        figure = block.figure
        _add_paragraph(doc, f"Диапазон глубин: {figure.depth_start:g}–{figure.depth_stop:g} м.")
        drawing = svg2rlg(BytesIO(figure.svg.encode("utf-8")))
        if drawing is None:
            raise RuntimeError("Не удалось преобразовать SVG-планшет для DOCX")
        png = renderPM.drawToString(drawing, fmt="PNG", dpi=220)
        stream = BytesIO(png)
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run()
        section = doc.sections[-1]
        usable_width = section.page_width - section.left_margin - section.right_margin
        run.add_picture(stream, width=usable_width)
        doc.add_paragraph()
        return
    figure = apply_report_plot_theme(block.figure)
    legend = _figure_report_legend(figure)
    depth_range = legend.get("depth_range", {}) if isinstance(legend.get("depth_range", {}), dict) else {}
    if depth_range:
        _add_paragraph(
            doc,
            f"Показан инженерно значимый диапазон глубин: "
            f"{float(depth_range.get('top', 0)):g}–{float(depth_range.get('base', 0)):g} м. "
            "Цветные зоны обозначают вероятный тип флюида; маркеры показывают кровлю, подошву и приоритетный интервал.",
        )
    intervals = list(legend.get("intervals", []) or [])
    if str(legend.get("report_kind", "")) == "detail" and intervals:
        table = doc.add_table(rows=1, cols=5)
        table.style = "Table Grid"
        headers = ("Интервал", "Глубина, м", "Мощность, м", "Флюид", "Достоверность")
        for idx, value in enumerate(headers):
            table.rows[0].cells[idx].text = value
            for run in table.rows[0].cells[idx].paragraphs[0].runs:
                run.bold = True
                run.font.size = Pt(8)
        for item in intervals:
            cells = table.add_row().cells
            values = (
                str(item.get("id", "")),
                f"{float(item.get('top', 0)):g}–{float(item.get('base', 0)):g}",
                f"{float(item.get('thickness', 0)):g}",
                str(item.get("fluid", "")),
                f"{float(item.get('confidence', 0)):g}%",
            )
            for idx, value in enumerate(values):
                cells[idx].text = value
                for run in cells[idx].paragraphs[0].runs:
                    run.font.size = Pt(8)
        doc.add_paragraph()
    # Curve names, interval colours and marker meanings are rendered directly
    # in the plot.  Repeated Unicode symbol tables caused black-square glyphs
    # in some Word/PDF font stacks and unnecessarily reduced plot scale.
    try:
        if hasattr(figure, "to_image"):
            png = figure.to_image(format="png", width=3200, height=2200, scale=1)
        elif hasattr(figure, "write_image"):
            buffer = BytesIO()
            figure.write_image(buffer, format="png", width=3200, height=2200)
            png = buffer.getvalue()
        else:
            raise TypeError("Figure backend does not support raster export")
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.add_run().add_picture(BytesIO(png), width=Inches(7.45))
        if str(legend.get("report_kind", "")) != "detail":
            _add_statistics_table(doc, list(legend.get("statistics", []) or []))
    except Exception as exc:
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(
            "График не удалось встроить в DOCX. Проверьте установку совместимой версии Kaleido "
            f"({type(exc).__name__})."
        )
        run.italic = True
    doc.add_paragraph()


def render_engineering_document_docx(
    document: EngineeringDocument,
    *,
    options: PresentationDocxOptions | None = None,
    on_progress: Callable[[int, str], None] | None = None,
    check_cancelled: Callable[[], None] | None = None,
) -> PresentationDocxResult:
    """Render a renderer-neutral EngineeringDocument into DOCX bytes."""

    def _check() -> None:
        if check_cancelled is not None:
            check_cancelled()

    def _progress(value: int, message: str) -> None:
        if on_progress is not None:
            on_progress(value, message)
        _check()

    _progress(2, "Инициализация DOCX")
    opts = options or PresentationDocxOptions()
    ensure_docx_available()
    doc = Document()
    _configure_document(doc, opts)

    title = document.metadata.title or opts.title
    _add_paragraph(doc, title, style="Title")
    if document.metadata.subtitle:
        _add_paragraph(doc, document.metadata.subtitle)
    _add_metadata_table(doc, document.metadata.rows)
    for note in document.metadata.notes:
        _add_paragraph(doc, note)

    section_total = max(1, len(document.sections))
    for section_index, section in enumerate(document.sections):
        _progress(10 + int((section_index / section_total) * 70), f"DOCX: раздел {section_index + 1} из {section_total}")
        if section.page_break_before and section_index > 0:
            doc.add_page_break()
        if section.title and (section_index > 0 or section.title not in {"Инженерные разделы отчета", "Разделы экспертного отчета"}):
            _add_paragraph(doc, section.title, style="Heading 1")
        for block in section.blocks:
            _check()
            if isinstance(block, DocumentTable):
                _add_document_table(doc, block)
            elif isinstance(block, DocumentNotice):
                _add_notice(doc, block)
            elif isinstance(block, DocumentPlot):
                _add_plot_placeholder(doc, block)
            elif isinstance(block, DocumentVisualizationPreview):
                _add_visualization_preview_placeholder(doc, block)

    _progress(84, "Сохранение DOCX")
    buffer = BytesIO()
    doc.save(buffer)
    _progress(98, "DOCX сформирован")
    return PresentationDocxResult(
        content=buffer.getvalue(),
        profile=document.metadata.profile,
        table_titles=document.table_titles,
        figure_count=document.plot_count + document.visualization_preview_count,
    )


def build_presentation_docx_report(
    model: PresentationModel,
    *,
    options: PresentationDocxOptions | None = None,
) -> PresentationDocxResult:
    """Render a DOCX report from PresentationModel through EngineeringDocument."""

    opts = options or PresentationDocxOptions()
    include_technical = opts.include_technical_appendix or model.metadata.report_profile == "expert"
    document = build_engineering_document(
        model,
        include_figures=opts.include_figures,
        include_technical_appendix=True if include_technical else None,
    )
    return render_engineering_document_docx(document, options=opts)


__all__ = [
    "PresentationDocxOptions",
    "PresentationDocxResult",
    "DOCX_AVAILABLE",
    "ensure_docx_available",
    "build_presentation_docx_report",
    "render_engineering_document_docx",
]
